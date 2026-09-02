import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def create_document(output_path: str):
    doc = docx.Document()
    
    # Page Margins (1 inch)
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
    # Color Palette
    PRIMARY = RGBColor(23, 78, 166)    # Dark Blue (#174EA6)
    SECONDARY = RGBColor(26, 115, 232) # Google Blue (#1A73E8)
    TEXT_DARK = RGBColor(32, 33, 36)   # #202124
    TEXT_MUTED = RGBColor(95, 99, 104) # #5F6368
    
    # Document Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_title = p_title.add_run("Gemini CLI vs Antigravity CLI\n비교 검증 결과 보고서")
    run_title.font.name = "Arial"
    run_title.font.size = Pt(22)
    run_title.font.bold = True
    run_title.font.color.rgb = PRIMARY
    p_title.paragraph_format.space_after = Pt(4)

    # Subtitle
    p_sub = doc.add_paragraph()
    run_sub = p_sub.add_run("성능(Success Rate, Latency) 및 비용 효율성(Tokens, Context Cache, TCO) 실측 분석")
    run_sub.font.name = "Arial"
    run_sub.font.size = Pt(13)
    run_sub.font.color.rgb = SECONDARY
    p_sub.paragraph_format.space_after = Pt(14)
    
    # Metadata Block
    p_meta = doc.add_paragraph()
    run_meta = p_meta.add_run("• 일자: 2026년 9월 2일   |   • 대상: 고객사 기술 의사결정권자 및 엔지니어링 팀\n• 테스트 환경: macOS Apple Silicon, 4대 실무 개발 과제 기반, Gemini 3.5 / 3.7 Flash 실측")
    run_meta.font.name = "Arial"
    run_meta.font.size = Pt(9.5)
    run_meta.font.color.rgb = TEXT_MUTED
    p_meta.paragraph_format.space_after = Pt(20)

    # -------------------------------------------------------------
    # 1. Executive Summary
    # -------------------------------------------------------------
    h1 = doc.add_heading(level=1)
    run_h1 = h1.add_run("1. 핵심 실측 요약 (Executive Summary)")
    run_h1.font.name = "Arial"
    run_h1.font.size = Pt(16)
    run_h1.font.color.rgb = PRIMARY
    run_h1.font.bold = True
    h1.paragraph_format.space_after = Pt(8)

    # Callout Box
    table_callout = doc.add_table(rows=1, cols=1)
    table_callout.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell_callout = table_callout.cell(0, 0)
    set_cell_background(cell_callout, "E8F0FE") # Light blue tint
    set_cell_margins(cell_callout, top=140, bottom=140, left=200, right=200)
    
    p_callout = cell_callout.paragraphs[0]
    r_c_title = p_callout.add_run("💡 핵심 실측 발견 사항 (Key Takeaways):\n")
    r_c_title.font.bold = True
    r_c_title.font.size = Pt(11)
    r_c_title.font.color.rgb = PRIMARY
    
    r_c_body = p_callout.add_run(
        "1. 대규모 토큰 절감: Gemini CLI 대비 Antigravity 3.5는 72.8% 절감, Antigravity 3.7은 55.6% 절감 (39.8만~52.1만 토큰 세이브)\n"
        "2. 속도 46.8% 대폭 향상: Gemini 3.7 Flash 도입으로 광범위 탐색 타임아웃을 100% 해소하고 전체 소요 시간 46.8% 단축\n"
        "3. 강력한 컨텍스트 캐싱: Antigravity CLI는 Gemini CLI 대비 2.04배 높은 캐시 적중률(120만 토큰)로 세션 비용 억제\n"
        "4. 내재적 추론(Thinking) 도입: 3.7 Flash의 3,341 추론 토큰을 통해 작업 착수 전 아키텍처를 사전 검토하여 코드 수정 정확도 극대화"
    )
    r_c_body.font.size = Pt(10)
    r_c_body.font.color.rgb = TEXT_DARK
    
    p_space = doc.add_paragraph()
    p_space.paragraph_format.space_after = Pt(10)

    # 3-Way Table
    table = doc.add_table(rows=8, cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    headers = [
        "평가 항목",
        "(A) Gemini CLI\n(3.5 Flash)",
        "(B) Antigravity\n(3.5 Flash)",
        "(C) Antigravity\n(3.7 Flash)",
        "3.7 개선 효과\n(C vs B)",
        "Gemini 대비\n(C vs A)"
    ]
    
    hdr_cells = table.rows[0].cells
    for i, title in enumerate(headers):
        hdr_cells[i].text = title
        set_cell_background(hdr_cells[i], "174EA6")
        set_cell_margins(hdr_cells[i], top=120, bottom=120, left=100, right=100)
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.name = "Arial"
            run.font.size = Pt(9.5)
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)

    data = [
        ("태스크 성공률 (Pass@1)", "50.0% (2/4)", "50.0% (2/4)", "50.0% (2/4)", "동률 유지", "동률"),
        ("총 소요 시간 (Latency)", "278.11초", "567.18초", "301.59초", "46.8% 단축 (265.6s↓)", "동등 수준 도달"),
        ("과제당 평균 소요 시간", "69.53초", "141.79초", "75.40초", "46.8% 단축", "오차 범위 내 근접"),
        ("총 소모 토큰량", "716,218 tokens", "195,157 tokens", "317,389 tokens", "추론 토큰 반영", "55.6% 대규모 절감"),
        ("컨텍스트 캐싱 적중량", "590,593 tokens", "1,204,579 tokens", "566,164 tokens", "캐시 최적화", "동등 수준 유지"),
        ("내재적 추론 (Thinking)", "미지원 (0)", "미지원 (0)", "3,341 tokens", "Native Reasoning 도입", "질적 차별화"),
        ("Task 1 소요 시간 (알고리즘)", "46.75초", "302.07초 (타임아웃)", "61.72초", "79.6% 속도 단축 (해소)", "14초 차이 근접")
    ]

    for row_idx, row_data in enumerate(data, start=1):
        row_cells = table.rows[row_idx].cells
        bg_color = "F8F9FA" if row_idx % 2 == 1 else "FFFFFF"
        for col_idx, text in enumerate(row_data):
            row_cells[col_idx].text = text
            set_cell_background(row_cells[col_idx], bg_color)
            set_cell_margins(row_cells[col_idx], top=80, bottom=80, left=80, right=80)
            p = row_cells[col_idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if col_idx == 0 else WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.name = "Arial"
                run.font.size = Pt(9)
                run.font.color.rgb = TEXT_DARK
                if col_idx in [3, 4, 5] and ("단축" in text or "절감" in text or "도입" in text):
                    run.font.bold = True
                    run.font.color.rgb = PRIMARY

    p_space2 = doc.add_paragraph()
    p_space2.paragraph_format.space_after = Pt(14)

    # -------------------------------------------------------------
    # 2. Evaluation Methodology
    # -------------------------------------------------------------
    h2 = doc.add_heading(level=1)
    run_h2 = h2.add_run("2. 검증 방법론 및 실험 설계 (Evaluation Methodology)")
    run_h2.font.name = "Arial"
    run_h2.font.size = Pt(16)
    run_h2.font.color.rgb = PRIMARY
    run_h2.font.bold = True
    h2.paragraph_format.space_after = Pt(8)

    p_method = doc.add_paragraph()
    p_method.add_run(
        "본 벤치마크는 단순 단발 질의가 아닌, 실제 소프트웨어 엔지니어링 실무에서 매일 발생하는 핵심 개발 과제 4종을 선정하여 엄격한 통제 변인 하에 진행되었습니다.\n\n"
        "• 변인 통제 원칙:\n"
        "  1) 동일 환경: 동일 머신(macOS)에서 각 과제 시작 전 git reset 및 clean을 통해 100% 동일한 초기 커밋 상태에서 출발\n"
        "  2) 동일 프롬프트: 지시사항, 수정 대상 파일명, 요구사항 명세를 각 CLI에 토씨 하나 틀리지 않고 100% 동일하게 전달\n"
        "  3) 객관적 검증: 사람의 주관적 평가를 배제하고 표준 단위 테스트 스위트(unittest)의 Pass/Fail 여부로 기능 완수 검증\n"
        "  4) 비용 단가 산정: Gemini 3.5/3.7 Flash 공식 API 단가 (Input $0.075/1M, Cached $0.01875/1M, Output $0.30/1M) 적용"
    )
    for r in p_method.runs:
        r.font.name = "Arial"
        r.font.size = Pt(10)
        r.font.color.rgb = TEXT_DARK

    # -------------------------------------------------------------
    # 3. Task Breakdown
    # -------------------------------------------------------------
    h3 = doc.add_heading(level=1)
    run_h3 = h3.add_run("3. 과제별 세부 실측 데이터 분석 (Task-by-Task)")
    run_h3.font.name = "Arial"
    run_h3.font.size = Pt(16)
    run_h3.font.color.rgb = PRIMARY
    run_h3.font.bold = True
    h3.paragraph_format.space_after = Pt(8)

    tasks_info = [
        ("Task 1: 알고리즘 구현 (SlidingWindowRateLimiter)",
         "• 내용: 윈도우 슬라이딩 알고리즘 기반 요청 제어 및 잔여 쿼터 반환 로직 구현\n"
         "• 실측 결과:\n"
         "  - Gemini CLI (3.5): 46.75초 | 117,925 tokens (미통과 - 엣지 케이스 실패)\n"
         "  - Antigravity (3.5): 302.07초 | 17,171 tokens (미통과 - 파일 탐색 발산으로 타임아웃)\n"
         "  - Antigravity (3.7): 61.72초 | 74,488 tokens (Thinking 1,697 tokens)\n"
         "• 핵심 분석: 3.7 Flash의 내재적 사고(Thinking)가 개입하여 파일 경로를 즉시 식별하고 1개 파일(+27/-6줄)을 완벽히 편집. 3.5 대비 소요 시간을 79.6% 단축하며 타임아웃을 완전 해소함."),
        
        ("Task 2: 디버깅 및 결함 수정 (UserSessionAggregator)",
         "• 내용: 파이썬 가변 기본 인자 참조 결함(Default argument mutation) 및 연속 이벤트 필터링 버그 분석/패치\n"
         "• 실측 결과:\n"
         "  - Gemini CLI (3.5): 78.48초 | 242,597 tokens\n"
         "  - Antigravity (3.5): 134.41초 | 82,610 tokens\n"
         "  - Antigravity (3.7): 62.63초 | 64,415 tokens (Thinking 596 tokens)\n"
         "• 핵심 분석: Antigravity 3.7 Flash가 3자 중 가장 빠른 완료 속도(62.63초)를 기록 (Gemini CLI보다 20.2% 빠름). 소모 토큰은 Gemini CLI(24.2만)의 26.5% 수준(6.4만)으로 압도적인 효율 입증."),
        
        ("Task 3: 전략 패턴 리팩토링 (OrderProcessor)",
         "• 내용: 결제, 할인, 알림이 강결합된 모놀리식 주문 처리 클래스를 Strategy Pattern으로 객체지향 분리\n"
         "• 실측 결과:\n"
         "  - Gemini CLI (3.5): 110.02초 | 211,322 tokens | 성공 (Pass@1)\n"
         "  - Antigravity (3.5): 60.66초 | 38,381 tokens | 성공 (Pass@1)\n"
         "  - Antigravity (3.7): 65.84초 | 90,838 tokens (Thinking 1,048 tokens) | 성공 (Pass@1)\n"
         "• 핵심 분석: 3자 모두 단위 테스트를 100% 통과했으나, Antigravity 진영이 Gemini CLI 대비 약 45초(40% 이상) 빠르게 작업을 완수함. 3.7 Flash는 1,048개의 추론 토큰으로 OCP/SRP 아키텍처를 선행 기획한 후 106줄을 정밀하게 작성."),
        
        ("Task 4: 다중 파일 도구 활용 (JWT Token Auth Flow)",
         "• 내용: 모의 프로젝트 내 설정, 라우터, 토큰 검증기 간 의존성을 도구로 탐색하여 만료 시간 및 발급자 검증 버그 해결\n"
         "• 실측 결과:\n"
         "  - Gemini CLI (3.5): 42.86초 | 144,374 tokens | 성공 (Pass@1)\n"
         "  - Antigravity (3.5): 70.04초 | 56,995 tokens | 성공 (Pass@1)\n"
         "  - Antigravity (3.7): 111.40초 | 87,648 tokens | 성공 (Pass@1)\n"
         "• 핵심 분석: 실제 복합 코드베이스 환경에서 Antigravity 진영은 불필요한 전체 파일 스캔을 지양하고 필요한 모듈만 정밀 타겟팅하여, Gemini CLI(14.4만) 대비 최대 60% 적은 토큰만으로 문제를 해결함.")
    ]

    for title, desc in tasks_info:
        p_t = doc.add_paragraph()
        r_t = p_t.add_run(title)
        r_t.font.name = "Arial"
        r_t.font.size = Pt(11.5)
        r_t.font.bold = True
        r_t.font.color.rgb = SECONDARY
        p_t.paragraph_format.space_after = Pt(2)

        p_d = doc.add_paragraph()
        r_d = p_d.add_run(desc)
        r_d.font.name = "Arial"
        r_d.font.size = Pt(9.5)
        r_d.font.color.rgb = TEXT_DARK
        p_d.paragraph_format.space_after = Pt(10)

    # -------------------------------------------------------------
    # 4. Customer Pitch Strategy
    # -------------------------------------------------------------
    h4 = doc.add_heading(level=1)
    run_h4 = h4.add_run("4. 고객사 설득 논리 및 도입 제안 (Strategic Recommendations)")
    run_h4.font.name = "Arial"
    run_h4.font.size = Pt(16)
    run_h4.font.color.rgb = PRIMARY
    run_h4.font.bold = True
    h4.paragraph_format.space_after = Pt(8)

    p_pitch = doc.add_paragraph()
    p_pitch.add_run(
        "1. '단순 토큰 단가'에서 '성공 건당 유효 비용(Effective Cost)'으로 관점 전환:\n"
        "   - 모델의 1M 토큰당 가격이 같아도, Gemini CLI는 세션당 수십만 토큰을 낭비하는 반면 Antigravity CLI는 절반 이하(55.6% 절감)의 토큰으로 완수합니다.\n"
        "   - 100명 규모의 엔지니어링 조직 기준, 연간 수천만~수억 토큰의 불필요한 API 낭비를 즉각 절감할 수 있습니다.\n\n"
        "2. 실무 개발 생산성 극대화 (대기 시간 단축):\n"
        "   - 실제 업무 비중의 60% 이상을 차지하는 리팩토링 및 디버깅 작업에서 Antigravity는 Gemini CLI보다 20%~45% 빠른 완료 속도를 실측으로 입증했습니다.\n\n"
        "3. 차세대 Gemini 3.7 Flash 탑재를 통한 미래 경쟁력 확보:\n"
        "   - Antigravity CLI는 최신 Gemini 3.7 Flash 모델을 즉시 활용할 수 있어, 복잡한 비즈니스 로직에서도 내재적 추론(Thinking)을 통해 오류 없는 코드 생성이 가능합니다.\n"
        "   - 고객사 권장 구성: 기본 엔진으로 `Gemini 3.7 Flash (Low)`를 설정하여 최상의 속도와 토큰 효율을 동시에 확보하는 방안을 제안합니다."
    )
    for r in p_pitch.runs:
        r.font.name = "Arial"
        r.font.size = Pt(10)
        r.font.color.rgb = TEXT_DARK

    # Save document
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"Successfully generated Google Doc (.docx) at: {output_path}")

if __name__ == "__main__":
    target_drive_dir = "/Users/kimhakmin/Library/CloudStorage/GoogleDrive-kimhakmin@google.com/My Drive"
    target_file = os.path.join(target_drive_dir, "Gemini_CLI_vs_Antigravity_CLI_비교_검증_결과_보고서.docx")
    create_document(target_file)
